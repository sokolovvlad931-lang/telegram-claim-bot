import asyncio
import logging
import os
from datetime import datetime
from io import BytesIO

from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile
from docx import Document
from docx.shared import Pt

# --- КОНФИГУРАЦИЯ ---
TOKEN = "8464321282:AAGySns6r_QESAV4TVrGh5mfQMfImVU3V_M"

# Настройка логирования
logging.basicConfig(level=logging.INFO)

# Инициализация бота и диспетчера
bot = Bot(token=TOKEN)
dp = Dispatcher()

# --- СОСТОЯНИЯ (FSM) ---
class ClaimStates(StatesGroup):
    choosing_marketplace = State()  # Выбор площадки
    entering_reason = State()       # Причина претензии
    entering_full_name = State()    # ФИО пользователя
    entering_address = State()     # Адрес для ответа
    entering_order_num = State()    # Номер заказа
    entering_price = State()        # Сумма претензии
    waiting_for_receipt = State()   # Ожидание фото чека (OCR)

# --- ЮРИДИЧЕСКИЕ СПРАВОЧНИКИ ---
LEGAL_BASE = {
    "WB": "ООО «Вайлдберриз», ИНН 7733545428, ОГРН 1067746062411. Юридический адрес: 142181, Московская область, г. Подольск, деревня Коледино, д. 6, стр. 1.",
    "OZON": "ООО «Интернет Решения», ИНН 7704217370, ОГРН 1027739244741. Юридический адрес: 123112, г. Москва, Пресненская наб., д. 10, блок С, эт. 41.",
    "Yandex": "ООО «ЯНДЕКС», ИНН 7736207543, ОГРН 1027700229193. Юридический адрес: 119021, г. Москва, ул. Льва Толстого, д. 16."
}

# --- КЛАВИАТУРЫ ---
def get_main_menu():
    buttons = [
        [InlineKeyboardButton(text="📝 Создать претензию", callback_data="create_claim")],
        [InlineKeyboardButton(text="📚 Правовой справочник", callback_data="legal_info")],
        [InlineKeyboardButton(text="📸 Распознать чек", callback_data="ocr_scan")]
    ]
    return InlineKeyboardMarkup(inline_keyboard=buttons)

def get_marketplaces():
    buttons = [
        [InlineKeyboardButton(text="Wildberries", callback_data="m_WB")],
        [InlineKeyboardButton(text="Ozon", callback_data="m_OZON")],
        [InlineKeyboardButton(text="Яндекс.Маркет", callback_data="m_Yandex")]
    ]
    return InlineKeyboardMarkup(inline_keyboard=buttons)

# --- ОБРАБОТЧИКИ (HANDLERS) ---

# Команда /start
@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    await message.answer(
        "👋 Добро пожаловать в **Юрист-Бот: Помощник по маркетплейсам**!\n\n"
        "Я помогу вам составить юридически грамотную претензию к Wildberries, Ozon или Яндекс.Маркет.\n"
        "Мои документы опираются на ГК РФ и Закон о защите прав потребителей (ЗоЗПП).",
        reply_markup=get_main_menu()
    )

# Обработка выбора создания претензии
@dp.callback_query(F.data == "create_claim")
async def start_claim_wizard(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.edit_text("Выберите маркетплейс, к которому у вас претензия:", reply_markup=get_marketplaces())
    await state.set_state(ClaimStates.choosing_marketplace)

# Выбор площадки
@dp.callback_query(ClaimStates.choosing_marketplace)
async def process_marketplace(callback: types.CallbackQuery, state: FSMContext):
    marketplace_id = callback.data.split("_")[1]
    await state.update_data(marketplace=marketplace_id)
    await callback.message.answer(f"Вы выбрали {marketplace_id}. Кратко опишите проблему (например: товар поврежден, задержка возврата денег, пришел не тот товар):")
    await state.set_state(ClaimStates.entering_reason)

# Ввод причины
@dp.message(ClaimStates.entering_reason)
async def process_reason(message: types.Message, state: FSMContext):
    await state.update_data(reason=message.text)
    await message.answer("Введите ваше ФИО полностью (например: Иванов Иван Иванович):")
    await state.set_state(ClaimStates.entering_full_name)

# Ввод ФИО
@dp.message(ClaimStates.entering_full_name)
async def process_name(message: types.Message, state: FSMContext):
    await state.update_data(full_name=message.text)
    await message.answer("Введите ваш почтовый адрес для получения ответа:")
    await state.set_state(ClaimStates.entering_address)

# Ввод адреса
@dp.message(ClaimStates.entering_address)
async def process_address(message: types.Message, state: FSMContext):
    await state.update_data(address=message.text)
    await message.answer("Введите номер заказа:")
    await state.set_state(ClaimStates.entering_order_num)

# Ввод номера заказа
@dp.message(ClaimStates.entering_order_num)
async def process_order(message: types.Message, state: FSMContext):
    await state.update_data(order_num=message.text)
    await message.answer("Введите общую сумму претензии (в рублях):")
    await state.set_state(ClaimStates.entering_price)

# Ввод цены и генерация DOCX
@dp.message(ClaimStates.entering_price)
async def process_price(message: types.Message, state: FSMContext):
    await state.update_data(price=message.text)
    data = await state.get_data()
    
    await message.answer("⏳ Формирую юридический документ по стандартам ГК РФ и ЗоЗПП...")
    
    # Генерация документа
    doc_stream = create_claim_docx(data)
    
    document_file = BufferedInputFile(
        doc_stream.getvalue(), 
        filename=f"Pretenziya_{data['marketplace']}.docx"
    )
    
    await message.answer_document(
        document_file, 
        caption="✅ Ваша претензия готова! Распечатайте её, подпишите и отправьте на юридический адрес площадки заказным письмом."
    )
    await state.clear()

# --- ФУНКЦИЯ ГЕНЕРАЦИИ DOCX ---
def create_claim_docx(data):
    doc = Document()
    
    # Стилизация
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Шапка документа
    header_info = LEGAL_BASE.get(data['marketplace'], "")
    p = doc.add_paragraph()
    p.add_run(f"Кому: {header_info}\n").bold = True
    p.add_run(f"От: {data['full_name']}\nАдрес: {data['address']}\n\n").bold = False
    
    # Заголовок
    title = doc.add_paragraph("ДОСУДЕБНАЯ ПРЕТЕНЗИЯ")
    title.alignment = 1 # Center
    
    # Текст претензии
    body = doc.add_paragraph()
    body.add_run(f"Мною в интернет-магазине {data['marketplace']} был совершен заказ №{data['order_num']}. ")
    body.add_run(f"В процессе исполнения обязательств возникла следующая проблема: {data['reason']}. ")
    body.add_run(f"Стоимость товара/услуги составляет {data['price']} руб.\n\n")
    
    body.add_run("В соответствии со ст. 18 Закона РФ «О защите прав потребителей», потребитель в случае обнаружения в товаре недостатков имеет право требовать полного возмещения убытков. Согласно ст. 309 ГК РФ, обязательства должны исполняться надлежащим образом.\n\n")
    
    body.add_run("ТРЕБУЮ:\n").bold = True
    body.add_run(f"Вернуть денежные средства в размере {data['price']} руб. в течение 10 календарных дней.\n\n")
    
    body.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')} ________________ (подпись)")

    # Сохранение в поток
    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# --- ПРАВОВОЙ СПРАВОЧНИК ---
@dp.callback_query(F.data == "legal_info")
async def show_legal_base(callback: types.CallbackQuery):
    text = (
        "⚖️ **Правовая база:**\n\n"
        "1. **ст. 18 ЗоЗПП**: Права потребителя при обнаружении недостатков.\n"
        "2. **ст. 22 ЗоЗПП**: Сроки удовлетворения отдельных требований (10 дней).\n"
        "3. **ст. 309 ГК РФ**: Обязательства должны исполняться надлежащим образом.\n"
        "4. **ст. 450.1 ГК РФ**: Отказ от договора в одностороннем порядке.\n\n"
        "Помните: претензия — это обязательный этап перед подачей иска в суд."
    )
    await callback.message.answer(text, parse_mode="Markdown")
    await callback.answer()

# --- OCR (РАСПОЗНАВАНИЕ ЧЕКОВ) ---
@dp.callback_query(F.data == "ocr_scan")
async def start_ocr(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.answer("📸 Отправьте фотографию чека. Я постараюсь извлечь из неё номер заказа и дату.\n(В данной версии реализована имитация Vision API)")
    await state.set_state(ClaimStates.waiting_for_receipt)

@dp.message(ClaimStates.waiting_for_receipt, F.photo)
async def process_ocr(message: types.Message, state: FSMContext):
    # В реальном приложении здесь вызывается Google Vision API, Tesseract или EasyOCR
    # Для демонстрации имитируем распознавание
    await message.answer("🔍 Анализирую фото чека...")
    await asyncio.sleep(2)
    
    extracted_order = "654432188"
    extracted_date = datetime.now().strftime("%d.%m.%Y")
    
    await message.answer(
        f"✅ Распознавание завершено!\n"
        f"**Номер заказа:** {extracted_order}\n"
        f"**Дата:** {extracted_date}\n\n"
        "Использовать эти данные при создании претензии?",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Да, продолжить", callback_data="create_claim")]
        ])
    )
    await state.clear()

# --- ЗАПУСК БОТА ---
async def main():
    print("Бот запущен и готов к работе...")
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:

        print("Бот остановлен.")
        # ================== ЗАПУСК ==================

async def main():
    print("🤖 Бот запущен и слушает Telegram")
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())
